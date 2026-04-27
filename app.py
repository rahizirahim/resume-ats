from flask import Flask, render_template, request, send_file, jsonify
import os
import pdfplumber
from groq import Groq
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import uuid

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

from dotenv import load_dotenv
load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text

def optimize_resume(resume_text):
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {
                "role": "system",
                "content": """You are an ATS resume formatter. You MUST follow this EXACT format strictly:

NAME: [Full Name Here]
JOBTITLE: [Desired Job Title]
PHONE: [Phone Here]
EMAIL: [Email Here]
LOCATION: [Location Here]
LINKEDIN: [LinkedIn if available, else leave blank]
AVAILABILITY: [Availability Here]

CAREER SUMMARY
[Write professional summary here]

EDUCATION
[Degree] | [Institution] | [Start Date - End Date]
- [CGPA or relevant detail]

EXPERIENCE
[Job Title] | [Company] | [Location] | [Start Date - End Date]
- [Achievement or responsibility]
- [Achievement or responsibility]

EXPERTISE
[Skill 1] | [Skill 2] | [Skill 3] | [Skill 4]

TECHNICAL SKILLS
[Skill 1] | [Skill 2] | [Skill 3] | [Skill 4]

IMPORTANT RULES:
1. Start contact fields with NAME:, JOBTITLE:, PHONE:, EMAIL:, LOCATION:, LINKEDIN:, AVAILABILITY:
2. Section headers CAREER SUMMARY, EDUCATION, EXPERIENCE, EXPERTISE, TECHNICAL SKILLS must be on their own line in CAPS
3. Use - for bullet points
4. Return ONLY the resume, no extra text"""
            },
            {
                "role": "user",
                "content": f"Optimize this resume:\n\n{resume_text}"
            }
        ]
    )
    return response.choices[0].message.content

def add_horizontal_line(paragraph, color="000000", size="6"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_photo_placeholder(doc):
    # Add a text box placeholder for photo
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("[ PHOTO ]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.bold = True
    # Add border around placeholder
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        bd = OxmlElement(f'w:{side}')
        bd.set(qn('w:val'), 'single')
        bd.set(qn('w:sz'), '12')
        bd.set(qn('w:space'), '4')
        bd.set(qn('w:color'), 'AAAAAA')
        pBdr.append(bd)
    pPr.append(pBdr)

def save_as_docx(text, output_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    lines = text.strip().split('\n')
    section_headers = ['CAREER SUMMARY', 'EDUCATION', 'EXPERIENCE', 'EXPERTISE', 'TECHNICAL SKILLS', 'CERTIFICATIONS', 'PROJECTS', 'SKILLS']

    name = ""
    jobtitle = ""
    contacts = []

    # Extract header info
    for line in lines:
        line = line.strip()
        if line.startswith('NAME:'):
            name = line.replace('NAME:', '').strip()
        elif line.startswith('JOBTITLE:'):
            jobtitle = line.replace('JOBTITLE:', '').strip()
        elif line.startswith('PHONE:'):
            contacts.append(line.replace('PHONE:', '').strip())
        elif line.startswith('EMAIL:'):
            contacts.append(line.replace('EMAIL:', '').strip())
        elif line.startswith('LOCATION:'):
            contacts.append(line.replace('LOCATION:', '').strip())
        elif line.startswith('LINKEDIN:'):
            val = line.replace('LINKEDIN:', '').strip()
            if val:
                contacts.append(val)
        elif line.startswith('AVAILABILITY:'):
            contacts.append(line.replace('AVAILABILITY:', '').strip())

    # Add photo placeholder
    add_photo_placeholder(doc)

    # Name + Job Title row
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(name)
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0, 0, 0)

    if jobtitle:
        name_para.add_run("    ")
        jt_run = name_para.add_run(jobtitle)
        jt_run.font.size = Pt(11)
        jt_run.font.color.rgb = RGBColor(80, 80, 80)
        jt_run.font.italic = True

    # Contact line
    contact_para = doc.add_paragraph()
    contact_text = "  |  ".join(contacts)
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = RGBColor(80, 80, 80)
    add_horizontal_line(contact_para)

    # Process rest of content
    skip_fields = ['NAME:', 'JOBTITLE:', 'PHONE:', 'EMAIL:', 'LOCATION:', 'LINKEDIN:', 'AVAILABILITY:']

    for line in lines:
        line = line.strip()
        if not line:
            continue
        if any(line.startswith(f) for f in skip_fields):
            continue

        # Section headers
        if line in section_headers:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            add_horizontal_line(p)

        # Bullet points
        elif line.startswith('-'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(f"• {line[1:].strip()}")
            run.font.size = Pt(10)

        # Job/Education title lines (contain |)
        elif '|' in line and not line.startswith('•'):
            parts = line.split('|')
            p = doc.add_paragraph()
            # First part bold (job title / degree)
            bold_run = p.add_run(parts[0].strip())
            bold_run.font.bold = True
            bold_run.font.size = Pt(10)
            # Rest normal
            for part in parts[1:]:
                normal_run = p.add_run(f"  |  {part.strip()}")
                normal_run.font.size = Pt(10)
                normal_run.font.color.rgb = RGBColor(80, 80, 80)

        # Normal text
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(10)

    doc.save(output_path)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/optimize', methods=['POST'])
def optimize():
    if 'resume_file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['resume_file']

    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not file.filename.endswith('.pdf'):
        return jsonify({'error': 'Please upload a PDF file'}), 400

    unique_id = str(uuid.uuid4())
    pdf_path = os.path.join(UPLOAD_FOLDER, f'{unique_id}.pdf')
    docx_filename = 'ATS_' + file.filename.replace('.pdf', '.docx')
    docx_path = os.path.join(OUTPUT_FOLDER, f'{unique_id}.docx')

    file.save(pdf_path)

    try:
        resume_text = extract_text_from_pdf(pdf_path)
        if not resume_text.strip():
            return jsonify({'error': 'Could not extract text from PDF'}), 400

        optimized_text = optimize_resume(resume_text)
        save_as_docx(optimized_text, docx_path)

    except Exception as e:
        import traceback
        print("ERROR:", traceback.format_exc())
        return jsonify({'error': f'Optimization failed: {str(e)}'}), 500
    finally:
        if os.path.exists(pdf_path):
            os.remove(pdf_path)

    return send_file(
        docx_path,
        as_attachment=True,
        download_name=docx_filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(debug=True)